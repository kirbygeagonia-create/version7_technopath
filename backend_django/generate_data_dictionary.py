import os
import sys
import django

# Setup Django environment
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'technopath.settings')
django.setup()

from django.apps import apps

def get_model_fields(model):
    fields = []
    for f in model._meta.get_fields():
        if f.is_relation and f.many_to_many and not f.auto_created:
            continue  # Exclude reverse M2M relationships without a physical column
        
        # Only take physical fields that exist on THIS model's table
        if not hasattr(f, 'column') or not f.column:
            continue
        
        field_type_str = type(f).__name__
        
        is_pk = getattr(f, 'primary_key', False)
        is_fk = f.is_relation and f.many_to_one
        fk_table = f.related_model._meta.db_table if is_fk and f.related_model else ''
        
        fields.append({
            'name': f.column,
            'type': field_type_str,
            'null': getattr(f, 'null', False),
            'pk': is_pk,
            'fk': is_fk,
            'fk_table': fk_table,
            'description': getattr(f, 'help_text', '') or getattr(f, 'verbose_name', f.name).capitalize()
        })
    return fields

def create_data_dictionary():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print('Please install python-docx: pip install python-docx')
        sys.exit(1)

    document = Document()
    title = document.add_heading('Technopath Data Dictionary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    models = apps.get_models(include_auto_created=True)
    
    for model in models:
        table_name = model._meta.db_table
        model_name = model.__name__
        
        document.add_heading(f'Table: {table_name} ({model_name})', level=1)
        
        fields = get_model_fields(model)
        
        # Create table: Column | Type | Null/Not Null | PK/FK | Description
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Column Name', 'Data Type', 'Null/Not Null', 'PK / FK', 'Description']
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Bold headers
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                
        for field in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field['name']
            row_cells[1].text = field['type']
            row_cells[2].text = 'NULL' if field['null'] else 'NOT NULL'
            
            pk_fk_text = []
            if field['pk']: pk_fk_text.append('PK')
            if field['fk']: pk_fk_text.append(f"FK ({field['fk_table']})")
            row_cells[3].text = ', '.join(pk_fk_text) if pk_fk_text else '-'
            row_cells[4].text = str(field['description'])
            
        document.add_paragraph('---')

    # Add django_migrations manually 
    document.add_heading('Table: django_migrations (Migration History)', level=1)
    mig_table = document.add_table(rows=1, cols=5)
    mig_table.style = 'Table Grid'
    hdr_cells = mig_table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            
    mig_fields = [
        {'name': 'id', 'type': 'BigAutoField', 'null': False, 'pk': True, 'desc': 'Primary Key'},
        {'name': 'app', 'type': 'CharField', 'null': False, 'pk': False, 'desc': 'Django Application Name'},
        {'name': 'name', 'type': 'CharField', 'null': False, 'pk': False, 'desc': 'Migration Name'},
        {'name': 'applied', 'type': 'DateTimeField', 'null': False, 'pk': False, 'desc': 'Date/Time Applied'},
    ]
    for field in mig_fields:
        row_cells = mig_table.add_row().cells
        row_cells[0].text = field['name']
        row_cells[1].text = field['type']
        row_cells[2].text = 'NULL' if field['null'] else 'NOT NULL'
        row_cells[3].text = 'PK' if field['pk'] else '-'
        row_cells[4].text = field['desc']
        
    document.save('Technopath_Data_Dictionary_Complete.docx')
    print(f'Technopath_Data_Dictionary_Complete.docx successfully generated with {len(models) + 1} tables.')

if __name__ == '__main__':
    create_data_dictionary()