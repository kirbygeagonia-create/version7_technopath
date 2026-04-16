import docx
from docx import Document
from docx.shared import Pt
import json
import copy

def map_db_type(django_type, max_length=None):
    mapping = {
        'AutoField': 'int(11)',
        'BigAutoField': 'bigint(20)',
        'CharField': f"varchar({max_length if max_length else 255})",
        'TextField': 'longtext',
        'IntegerField': 'int(11)',
        'BigIntegerField': 'bigint(20)',
        'SmallIntegerField': 'smallint(6)',
        'PositiveIntegerField': 'int(10) unsigned',
        'PositiveSmallIntegerField': 'smallint(5) unsigned',
        'FloatField': 'double',
        'DecimalField': 'decimal(10,2)',
        'BooleanField': 'tinyint(1)',
        'NullBooleanField': 'tinyint(1)',
        'DateTimeField': 'datetime',
        'DateField': 'date',
        'TimeField': 'time',
        'FileField': 'varchar(255)',
        'ImageField': 'varchar(255)',
        'EmailField': f"varchar({max_length if max_length else 254})",
        'URLField': 'varchar(200)',
        'UUIDField': 'char(32)',
        'ForeignKey': 'int(11)',
        'OneToOneField': 'int(11)',
    }
    return mapping.get(django_type, 'varchar(255)')

with open('schema_dump_fixed.json', 'r') as f:
    schema = json.load(f)

doc_path = r'c:\Users\ADMIN\OneDrive\Documents\SAD_System_files\technopath\technopath_pwa\version4_technopath\Technopath -Chapter 1 and Chapter 3 -  Final(2) (Repaired).docx'
doc = Document(doc_path)
scratch = Document()

def replace_element_with_new_elements(old_element, new_elements):
    parent = old_element.getparent()
    if parent is None:
        print("Cannot replace element, no parent found.")
        return
    for new_el in new_elements:
        old_element.addprevious(copy.deepcopy(new_el))
    parent.remove(old_element)

def make_data_dict_table(table_name, info):
    t = scratch.add_table(rows=0, cols=5)
    t.style = 'Table Grid'
    
    r1 = t.add_row()
    r1.cells[0].text = "Table Name:"
    r1.cells[1].text = table_name
    for i in range(1, 4): r1.cells[1].merge(r1.cells[i+1])
    r1.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r2 = t.add_row()
    r2.cells[0].text = "Description:"
    r2.cells[1].text = f"Stores records for {info['model_name']}."
    for i in range(1, 4): r2.cells[1].merge(r2.cells[i+1])
    r2.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r3 = t.add_row()
    headers = ["Attribute Name", "Data Type", "Default Value", "Null", "Description"]
    for i, h in enumerate(headers):
        r3.cells[i].text = h
        r3.cells[i].paragraphs[0].runs[0].font.bold = True
        
    for field in info['fields']:
        if not field['name']: continue
        tr = t.add_row()
        tr.cells[0].text = field['name']
        
        dtype = map_db_type(field['type'], field.get('max_length'))
        tr.cells[1].text = dtype
        
        # Determine Default Value / PK
        def_val = "PK" if field['pk'] else field.get('default')
        if def_val is None: def_val = "-"
        tr.cells[2].text = str(def_val)
        
        tr.cells[3].text = "Yes" if field['null'] else "No"
        desc = str(field['help_text']) if field['help_text'] else str(field['verbose_name'])
        if field['fk']: desc += f" (Foreign Key to {field['fk_table']})"
        tr.cells[4].text = desc
        
    return t._element

def make_use_case_narrative(row_data):
    # row_data format: 0:Name, 1:Actor, 2:Desc, 3:Precond, 4:Dyn Precond, 5:Trigger, 6:Flow
    t = scratch.add_table(rows=7, cols=2)
    t.style = 'Table Grid'
    
    r1 = t.rows[0]
    r1.cells[0].text = f"Use Case Name:  {row_data[0]}"
    r1.cells[1].text = "Technopath"
    r1.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r2 = t.rows[1]
    r2.cells[0].text = f"Primary Actor: {row_data[1]}"
    r2.cells[1].text = f"Primary Actor: {row_data[1]}"
    r2.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r3 = t.rows[2]
    r3.cells[0].text = f"Description:  {row_data[2]}"
    r3.cells[1].text = f"Description:  {row_data[2]}"
    r3.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r4 = t.rows[3]
    r4.cells[0].text = f"Preconditions: {row_data[3]}"
    r4.cells[1].text = f"Preconditions: {row_data[3]}"
    r4.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r5 = t.rows[4]
    r5.cells[0].text = f"Dynamic Precondition:  {row_data[4]}"
    r5.cells[1].text = f"Dynamic Precondition:  {row_data[4]}"
    r5.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r6 = t.rows[5]
    r6.cells[0].text = f"Trigger / Event:  {row_data[5]}"
    r6.cells[1].text = f"Trigger / Event:  {row_data[5]}"
    r6.cells[0].paragraphs[0].runs[0].font.bold = True
    
    r7 = t.rows[6]
    r7.cells[0].text = f"Basic Flow:  {row_data[6]}\nType:\tExternal\tTemporary"
    r7.cells[1].text = f"Basic Flow:  {row_data[6]}\nType:\tExternal\tTemporary"
    r7.cells[0].paragraphs[0].runs[0].font.bold = True
    
    return t._element

# Strategy: Collect tables to replace
table_actions = []
use_case_data = []

for i, t in enumerate(doc.tables):
    try:
        header = t.rows[0].cells[0].text.strip()
        if header == 'Use Case Name':  
            # Extract 10 rows (skip header)
            for r in t.rows[1:]:
                use_case_data.append([c.text.strip() for c in r.cells])
            table_actions.append((t._element, "USE_CASE"))
        elif header == 'Column Name':  
            table_actions.append((t._element, "DATA_DICT"))
        elif header == 'Component' and i == 32:  
            table_actions.append((t._element, "SPEC_PROGRAM"))
        elif header == 'Component' and i == 33:  
            table_actions.append((t._element, "SPEC_SOFTWARE"))
        elif header == 'Specification' and i == 34:  
            table_actions.append((t._element, "SPEC_HARDWARE"))
    except:
        pass

data_dict_inserted = False

for el, typ in table_actions:
    if typ == "USE_CASE":
        new_elements = []
        for uc in use_case_data:
            new_elements.append(make_use_case_narrative(uc))
            p = scratch.add_paragraph()
            new_elements.append(p._element)
        replace_element_with_new_elements(el, new_elements)
        
    elif typ == "DATA_DICT":
        if not data_dict_inserted:
            new_elements = []
            for t_name, t_info in schema.items():
                if t_name:
                    new_elements.append(make_data_dict_table(t_name, t_info))
                    p = scratch.add_paragraph()
                    new_elements.append(p._element)
            replace_element_with_new_elements(el, new_elements)
            data_dict_inserted = True
        else:
            parent = el.getparent()
            if parent is not None: parent.remove(el)

    elif typ == "SPEC_PROGRAM":
        new_elements = []
        p_hdr = scratch.add_paragraph()
        p_hdr.add_run("Program Specification").bold = True
        new_elements.append(p_hdr._element)
        new_elements.append(scratch.add_paragraph("Language: Python (Django & Flask), JavaScript (Vue.js)")._element)
        new_elements.append(scratch.add_paragraph("Database: Oracle Database / MySQL")._element)
        new_elements.append(scratch.add_paragraph("IDE: Visual Studio Code")._element)
        replace_element_with_new_elements(el, new_elements)
        
    elif typ == "SPEC_SOFTWARE":
        new_elements = []
        p_hdr = scratch.add_paragraph()
        p_hdr.add_run("Software Requirement").bold = True
        new_elements.append(p_hdr._element)
        new_elements.append(scratch.add_paragraph("Server OS: Windows Server / Linux Distribution")._element)
        new_elements.append(scratch.add_paragraph("Client Browser: Google Chrome, Safari, or Mozilla Firefox")._element)
        replace_element_with_new_elements(el, new_elements)
        
    elif typ == "SPEC_HARDWARE":
        new_elements = []
        p_hdr = scratch.add_paragraph()
        p_hdr.add_run("Hardware Requirement").bold = True
        new_elements.append(p_hdr._element)
        new_elements.append(scratch.add_paragraph("Server Equipment:")._element)
        new_elements.append(scratch.add_paragraph("- Processor: Quad-Core 2.4 GHz or higher", style='List Paragraph')._element)
        new_elements.append(scratch.add_paragraph("- Memory (RAM): Minimum 8GB", style='List Paragraph')._element)
        new_elements.append(scratch.add_paragraph("- Storage: 256GB SSD", style='List Paragraph')._element)
        new_elements.append(scratch.add_paragraph("Client Equipment:")._element)
        new_elements.append(scratch.add_paragraph("- Mobile Smartphone (Android/iOS) with camera and internet connectivity.", style='List Paragraph')._element)
        replace_element_with_new_elements(el, new_elements)

# CV Code remains unchanged
in_cv = False
for p in doc.paragraphs:
    if 'CURRICULUM VITAE' in p.text.upper():
        in_cv = True
        p.runs[0].bold = True
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    if in_cv:
        if len(p.text) > 0 and len(p.text) < 40 and not p.text.startswith(' '):
            for r in p.runs: r.bold = True
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

doc.save('Technopath_Document_Final_V2.docx')
print("Successfully generated revised Technopath_Document_Final_V2.docx!")
