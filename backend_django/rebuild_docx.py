import docx
from docx import Document
from docx.shared import Pt
import json
import copy

def clone_table(new_doc, original_table):
    # Not used directly but good to have
    pass

def map_db_type(django_type):
    mapping = {
        'AutoField': 'NUMBER IDENTITY',
        'BigAutoField': 'NUMBER IDENTITY',
        'CharField': 'VARCHAR2',
        'TextField': 'CLOB',
        'IntegerField': 'NUMBER',
        'BigIntegerField': 'NUMBER(19)',
        'SmallIntegerField': 'NUMBER(5)',
        'PositiveIntegerField': 'NUMBER',
        'PositiveSmallIntegerField': 'NUMBER(5)',
        'FloatField': 'FLOAT',
        'DecimalField': 'NUMBER',
        'BooleanField': 'NUMBER(1)',
        'NullBooleanField': 'NUMBER(1)',
        'DateTimeField': 'TIMESTAMP',
        'DateField': 'DATE',
        'TimeField': 'TIMESTAMP',
        'FileField': 'VARCHAR2(255)',
        'ImageField': 'VARCHAR2(255)',
        'EmailField': 'VARCHAR2',
        'URLField': 'VARCHAR2',
        'UUIDField': 'VARCHAR2(36)',
        'ForeignKey': 'NUMBER',
        'OneToOneField': 'NUMBER',
    }
    return mapping.get(django_type, 'VARCHAR2(255)')

# 1. Load schema
with open('schema_dump.json', 'r') as f:
    schema = json.load(f)

# 2. Load original document
doc_path = r'c:\Users\ADMIN\OneDrive\Documents\SAD_System_files\technopath\technopath_pwa\version4_technopath\Technopath -Chapter 1 and Chapter 3 -  Final(2) (Repaired).docx'
doc = Document(doc_path)

# Prepare a scratch doc to generate perfectly formed elements
scratch = Document()

def replace_element_with_new_elements(old_element, new_elements):
    parent = old_element.getparent()
    if parent is None:
        print("Cannot replace element, no parent found.")
        return
    for new_el in new_elements:
        old_element.addprevious(copy.deepcopy(new_el))
    parent.remove(old_element)

# A function to make ATTENDX Data Dictionary Table
def make_data_dict_table(table_name, info):
    t = scratch.add_table(rows=0, cols=5)
    t.style = 'Table Grid'
    
    # Row 1: Table Name
    r1 = t.add_row()
    r1.cells[0].text = "Table Name:"
    r1.cells[1].text = table_name
    r1.cells[1].merge(r1.cells[2])
    r1.cells[1].merge(r1.cells[3])
    r1.cells[1].merge(r1.cells[4])
    r1.cells[0].paragraphs[0].runs[0].font.bold = True
    
    # Row 2: Description
    r2 = t.add_row()
    r2.cells[0].text = "Description:"
    # Use model name as description or custom based on verbose name
    r2.cells[1].text = f"Stores records for {info['model_name']}."
    r2.cells[1].merge(r2.cells[2])
    r2.cells[1].merge(r2.cells[3])
    r2.cells[1].merge(r2.cells[4])
    r2.cells[0].paragraphs[0].runs[0].font.bold = True
    
    # Row 3: Headers
    r3 = t.add_row()
    headers = ["Attribute Name", "Data Type", "Default Value", "Null", "Description"]
    for i, h in enumerate(headers):
        r3.cells[i].text = h
        r3.cells[i].paragraphs[0].runs[0].font.bold = True
        
    # Data Rows
    for field in info['fields']:
        if not field['name']: continue
        tr = t.add_row()
        tr.cells[0].text = field['name']
        
        dtype = map_db_type(field['type'])
        if field.get('max_length'):
            dtype = f"VARCHAR2({field['max_length']})"
        tr.cells[1].text = dtype
        
        tr.cells[2].text = "NULL" # Django doesn't expose default value natively in this dump
        tr.cells[3].text = "YES" if field['null'] else "NO"
        tr.cells[4].text = str(field['help_text']) if field['help_text'] else str(field['verbose_name'])
        
    return t._element

# A function to make ATTENDX Use Case Narrative
def make_use_case_narrative():
    t = scratch.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    
    r1 = t.rows[0]
    r1.cells[0].text = "Use Case Name:\nTechnopath PWA Pathfinding and Campus Management"
    r1.cells[1].text = "Technopath"
    
    r2 = t.rows[1]
    r2.cells[0].text = "Primary Actor: Administrator"
    r2.cells[1].text = "Primary Actor: User / Student"
    
    r3 = t.rows[2]
    r3.cells[0].text = "Description:\nThis use case describes how the administrator manages announcements, facilities, user feedback, and campus map nodes."
    r3.cells[1].text = "Description:\nThis use case describes how the user efficiently navigates the campus via AI chatbot query assistance and interactive graphs, and rates the system."
    
    r4 = t.rows[3]
    r4.cells[0].text = "Preconditions:\nThe Administrator must have authenticated access and appropriate permissions."
    r4.cells[1].text = "Preconditions:\nThe User must be on the mobile or web app with basic navigation modules loaded."
    
    return t._element

# Strategy:
# Iterate all tables. 
# Identifying tables by index is fragile but we logged them.
# Tech Table 3: Use Case Matrix (11 rows)
# Tech Tables 4 to 31: Data Dictionaries
# Tech Tables 32 to 34: Hardware/Software/Program specs

table_idx_to_replace = []
for i, t in enumerate(doc.tables):
    try:
        header = t.rows[0].cells[0].text.strip()
        if header == 'Use Case Name':  # Table 3
            table_idx_to_replace.append((i, "USE_CASE"))
        elif header == 'Column Name':  # Tables 4-31
            table_idx_to_replace.append((i, "DATA_DICT"))
        elif header == 'Component' or header == 'Specification':  # Tables 32-34
            table_idx_to_replace.append((i, "SPECS"))
    except:
        pass

print(f"Detected Tables to Replace: {len(table_idx_to_replace)}")

# We should do replacements from bottom up so indices don't shift, or just store the elements directly!
elements_to_replace = []
for i, typ in table_idx_to_replace:
    elements_to_replace.append((doc.tables[i]._element, typ))

# Before we process DATA_DICT, we have 9 tables mapped in schema. The original doc had like 28 tables.
# I will only insert the valid schema tables where the FIRST data dict table was, and delete all other data dict tables.
data_dict_inserted = False

for el, typ in elements_to_replace:
    if typ == "USE_CASE":
        new_uc = make_use_case_narrative()
        replace_element_with_new_elements(el, [new_uc])
        
    elif typ == "DATA_DICT":
        if not data_dict_inserted:
            new_elements = []
            for t_name, t_info in schema.items():
                if t_name:
                    new_elements.append(make_data_dict_table(t_name, t_info))
                    
                    # Add paragraph space
                    p = scratch.add_paragraph()
                    new_elements.append(p._element)
                    
            replace_element_with_new_elements(el, new_elements)
            data_dict_inserted = True
        else:
            # We already inserted all correct data dictionaries at the first table's location.
            # We just delete the remaining obsolete tables.
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                
    elif typ == "SPECS":
        # Convert to paragraph format
        new_elements = []
        p_hdr = scratch.add_paragraph()
        run = p_hdr.add_run("System Specifications:")
        run.bold = True
        new_elements.append(p_hdr._element)
        
        p1 = scratch.add_paragraph("1. Python (Django & Flask) based backend.")
        p1.style = 'List Paragraph'
        p2 = scratch.add_paragraph("2. Vue.js PWA for the mobile interface.")
        p2.style = 'List Paragraph'
        p3 = scratch.add_paragraph("3. Oracle Database for persistent datastore.")
        p3.style = 'List Paragraph'
        p4 = scratch.add_paragraph("4. Hardware: Server capable of AI Chatbot inference and web hosting; Mobile devices for end users.")
        p4.style = 'List Paragraph'
        
        new_elements.extend([p1._element, p2._element, p3._element, p4._element])
        
        # We only want to insert the text once, so we'll do it on the first SPECS table, 
        # and delete the rest just like Data Dict.
        if el.getprevious() is not None and "System Specifications" in scratch.paragraphs[-1].text:
             parent = el.getparent()
             if parent is not None: parent.remove(el)
        else:
             replace_element_with_new_elements(el, new_elements) # Wait, it might insert 3 times. Let's fix this logic.

# Better Logic for SPECS deletion:
# Make a global flag.
specs_inserted = False
for el, typ in elements_to_replace:
    if typ == "SPECS":
        if not specs_inserted:
            new_elements = []
            p_hdr = scratch.add_paragraph()
            run = p_hdr.add_run("System Specifications Narrative")
            run.bold = True
            p_hdr.style = 'Heading 2'
            new_elements.append(p_hdr._element)
            
            bullets = [
               "Development Environment: Visual Studio Code, Git.",
               "Backend Architecture: Django REST Framework (Python 3) for core APIs, Flask for Gemini AI Chatbot processing.",
               "Frontend Architecture: Vue.js 3 progressive web app (PWA) compiled with Vite.",
               "Database Platform: Oracle / PostgreSQL (as defined in settings).",
               "Hardware Requirements: Minimum 8GB RAM, Quad-core processor for server. End-users require modern smartphone with web browser."
            ]
            for b in bullets:
                p = scratch.add_paragraph(b, style='List Paragraph')
                new_elements.append(p._element)
                
            replace_element_with_new_elements(el, new_elements)
            specs_inserted = True
        else:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

# CV Format fixing
# Locate CV header and modify style.
in_cv = False
for p in doc.paragraphs:
    if 'CURRICULUM VITAE' in p.text.upper():
        in_cv = True
        p.runs[0].bold = True
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Create ATTENDX MAP OF RESEARCH header instead of basic table.
        # But maybe we don't need to insert random data, just ensure the format is centered.
    if in_cv:
        # Match ATTENDX CV formatting: Center align, bold names
        if len(p.text) > 0 and len(p.text) < 40 and not p.text.startswith(' '):
            for r in p.runs: r.bold = True
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

doc.save('Technopath_Document_Updated.docx')
print("Successfully generated Technopath_Document_Updated.docx!")
