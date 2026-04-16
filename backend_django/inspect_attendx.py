import docx

doc_a = docx.Document(r'c:\Users\ADMIN\OneDrive\Documents\SAD_System_files\technopath\technopath_pwa\version4_technopath\ATTENDX-EXAMPLE-DOCUMENT.docx')

print("=== ATTENDX TABLE 4 (Use Case Matrix?) ===")
try:
    for row in doc_a.tables[4].rows:
        print(" | ".join([c.text.strip().replace('\n', ' ') for c in row.cells]))
except Exception as e:
    print(e)
    
print("\n=== ATTENDX TABLE 5 (Data Dictionary) ===")
try:
    for row in doc_a.tables[5].rows:
        print(" | ".join([c.text.strip().replace('\n', ' ') for c in row.cells]))
except Exception as e:
    print(e)
    
print("\n=== ATTENDX Hardware / Software Spec ===")
# Let's find "Hardware Requirement"
in_hw = False
for p in doc_a.paragraphs:
    text = p.text.strip()
    if 'Hardware Requirement' in text or 'Program Specification' in text or 'Software Requirement' in text:
        in_hw = True
        print(f"HEADER: {text}")
        continue
    if in_hw and text:
        print(text)
        if 'Use Case' in text or 'Data Dictionary' in text:
            in_hw = False
            
print("\n=== ATTENDX CV ===")
in_cv = False
for p in doc_a.paragraphs:
    text = p.text.strip()
    if 'CURRICULUM VITAE' in text:
        in_cv = True
        print(f"HEADER: {text}")
        continue
    if in_cv and text:
        print(text[:100]) # only print first 100 chars to avoid massive output
