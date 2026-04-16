import docx
import sys

doc_t = docx.Document(r'c:\Users\ADMIN\OneDrive\Documents\SAD_System_files\technopath\technopath_pwa\version4_technopath\Technopath -Chapter 1 and Chapter 3 -  Final(2) (Repaired).docx')
print('Technopath Tables:')
for i, t in enumerate(doc_t.tables):
    try:
        col2 = t.rows[0].cells[1].text[:30].replace('\n', ' ') if len(t.rows[0].cells) > 1 else ''
        print(f'Tech Table {i}: {len(t.rows)} rows. Header: {t.rows[0].cells[0].text[:30].replace(chr(10), " ")} | {col2}')
    except Exception as e:
        print(f'Tech Table {i}: Error {e}')

print('\nAttendx Tables:')
doc_a = docx.Document(r'c:\Users\ADMIN\OneDrive\Documents\SAD_System_files\technopath\technopath_pwa\version4_technopath\ATTENDX-EXAMPLE-DOCUMENT.docx')
for i, t in enumerate(doc_a.tables):
    try:
        col2 = t.rows[0].cells[1].text[:30].replace('\n', ' ') if len(t.rows[0].cells) > 1 else ''
        print(f'Attendx Table {i}: {len(t.rows)} rows. Header: {t.rows[0].cells[0].text[:30].replace(chr(10), " ")} | {col2}')
    except Exception as e:
        print(f'Attendx Table {i}: Error {e}')
