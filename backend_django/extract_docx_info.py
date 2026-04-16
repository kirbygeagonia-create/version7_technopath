import docx
import json
import sys

def dump_docx(filepath, outpath):
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"Error opening {filepath}: {e}")
        return

    data = []
    for element in doc.elements if hasattr(doc, 'elements') else doc.iter_inner_content():
        pass # Not easily available in basic docx

    # Basic approach: just get paragraphs and tables
    # Unfortunately ordering of paragraphs vs tables is lost if we just iterate doc.paragraphs then doc.tables
    # Let's use a workaround with doc._element.body
    
    # Actually python-docx doesn't provide an easy way to iterate paras and tables in order out of the box.
    # Let's just output paragraphs and tables separately for analysis.

if __name__ == "__main__":
    pass
